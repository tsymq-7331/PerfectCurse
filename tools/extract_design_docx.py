from docx import Document

path = r"C:\Users\唐宋元明清\Downloads\完美诅咒Mod_实现思路与执行文档.docx"
doc = Document(path)

for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if text:
        print(text)

for table_index, table in enumerate(doc.tables, 1):
    print(f"\n[TABLE {table_index}]")
    for row in table.rows:
        print("\t".join(cell.text.replace("\n", " / ").strip() for cell in row.cells))
